"""Word .docx → Markdown：mammoth 首选，失败时回落 python-docx。"""
from __future__ import annotations

import logging
from pathlib import Path

import mammoth

from .base import BaseConverter, ConversionResult

logger = logging.getLogger(__name__)


class DocxConverter(BaseConverter):
    kind = "docx"

    def convert(self, path: Path) -> ConversionResult:
        warnings: list[str] = []

        # 1) mammoth 首选：直接转 markdown
        try:
            with path.open("rb") as f:
                result = mammoth.convert_to_markdown(f)
            md = result.value or ""
            for msg in result.messages:
                # 只把 warning/error 级带上来
                lvl = getattr(msg, "type", "info")
                if lvl in ("warning", "error"):
                    warnings.append(f"mammoth-{lvl}: {msg.message}")

            if md.strip():
                if not md.endswith("\n"):
                    md += "\n"
                return ConversionResult(
                    status="success",
                    markdown=md,
                    warnings=warnings,
                    extra={"chars": len(md), "engine": "mammoth"},
                )
            # mammoth 出空 → 兜底
            warnings.append("mammoth 输出为空，尝试用 python-docx 兜底")
        except Exception as e:
            logger.debug(f"mammoth 失败 {path}: {e}")
            warnings.append(f"mammoth 失败（{e.__class__.__name__}），尝试 python-docx 兜底")

        # 2) python-docx 兜底
        try:
            md = _convert_with_python_docx(path)
            if not md.endswith("\n"):
                md += "\n"
            return ConversionResult(
                status="success",
                markdown=md,
                warnings=warnings,
                extra={"chars": len(md), "engine": "python-docx"},
            )
        except Exception as e:
            return ConversionResult(
                status="failed",
                markdown="",
                warnings=warnings,
                error=f"docx 转换失败: {e.__class__.__name__}: {e}",
            )


def _convert_with_python_docx(path: Path) -> str:
    """简单兜底：标题段落 + 普通段落 + 表格 → markdown。"""
    from docx import Document  # type: ignore

    doc = Document(str(path))
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.rstrip()
        if not text:
            lines.append("")
            continue
        style = (para.style.name or "").lower()
        if style.startswith("heading 1") or style == "title":
            lines.append(f"# {text}")
        elif style.startswith("heading 2"):
            lines.append(f"## {text}")
        elif style.startswith("heading 3"):
            lines.append(f"### {text}")
        elif style.startswith("heading 4"):
            lines.append(f"#### {text}")
        elif style.startswith("heading 5"):
            lines.append(f"##### {text}")
        elif style.startswith("heading 6"):
            lines.append(f"###### {text}")
        elif "list" in style:
            lines.append(f"- {text}")
        else:
            lines.append(text)

    # 表格
    for i, table in enumerate(doc.tables):
        lines.append("")
        lines.append(f"**Table {i + 1}**")
        rows = []
        for row in table.rows:
            cells = [cell.text.replace("\n", " ").replace("|", "\\|") for cell in row.cells]
            rows.append(cells)
        if not rows:
            continue
        ncols = max(len(r) for r in rows)
        # 补齐
        rows = [r + [""] * (ncols - len(r)) for r in rows]
        header = rows[0]
        body = rows[1:] if len(rows) > 1 else []
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join("---" for _ in header) + " |")
        for r in body:
            lines.append("| " + " | ".join(r) + " |")
        lines.append("")

    return "\n".join(lines)
